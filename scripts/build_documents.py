"""Build the submission documents: the combined supplementary PDF, DOCX and CSV renderings.

Scientific Reports asks for supplementary items combined into a single PDF, with the manuscript
title and author list on its first page, every supplementary figure carrying a title and legend
in the file, and every supplementary table carrying a title. Spreadsheets may stay separate.
This assembles that file, and produces Word and PDF renderings of the text and table sources so
the package can be read without a Markdown viewer.

    python scripts/build_documents.py --sources DIR --out DIR

Sources are the generated supplementary items; nothing here computes a result, so a rerun after
regenerating the tables is safe and idempotent.
"""

from __future__ import annotations

import argparse
import csv
import pathlib
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

TITLE = (
    "Quantum simulation of mutation-selection dynamics reveals a structural "
    "obstruction to quantum advantage"
)
AUTHORS = (
    "Anees Ahmed Mahaboob Ali, Everette Jacob Remington Nelson, "
    "Radhakrishnan Delhibabu"
)
AFFILIATION = (
    "Gene Therapy Laboratory, School of Bio Sciences and Technology, and School of Computer "
    "Science and Engineering, Vellore Institute of Technology, Vellore 632014, Tamil Nadu, India"
)

# Title and legend for each supplementary figure, embedded in the combined file as the journal
# requires. Panels are not cited individually, per the same instruction.
FIGURE_LEGENDS = {
    "S1_gap_and_conditioning.png": (
        "Supplementary Figure S1. Spectral gap at the error threshold and the conditioning "
        "that follows from it.",
        "Left: the minimum spectral gap at the critical mutation rate against system size, "
        "computed at arbitrary precision and on a 1,500-point grid. The grid reads the minimum "
        "as far larger than it is once the gap falls below its resolution, which is why the "
        "extended-precision path is required rather than preferred. Right: the worst "
        "eigenvector condition number over the same range. Conditioning degrades as the gap "
        "closes, which bounds the accuracy attainable by any eigenvector-extraction method, "
        "quantum or classical.",
    ),
    "S2_barren_plateau.png": (
        "Supplementary Figure S2. Gradient-variance decay for the variational route.",
        "Variance of the McLachlan force against system size for every landscape and statistic "
        "combination measured, on logarithmic axes. Fitted decay bases are given in the legend "
        "and lie between 0.535 and 0.556. Exponential decay of the gradient variance with "
        "system size is the barren-plateau behaviour expected of variational circuits, and it "
        "bounds the sizes at which the variational route remains trainable.",
    ),
    "S3_symmetry_breaking.png": (
        "Supplementary Figure S3. Finite-population symmetry breaking in spin-glass landscapes.",
        "Cosine similarity of the Wright-Fisher baseline against the exact distribution, by "
        "landscape family, one point per problem instance. Spin-glass fitness is invariant "
        "under simultaneous inversion of all spins, so the exact distribution is symmetric "
        "under complementation while a finite population settles into one branch of the pair. "
        "The cosine similarity between a one-branch and a symmetric two-branch distribution is "
        "exactly the value marked, and the spin-glass instances cluster there. The effect is a "
        "property of finite populations rather than a defect of the implementation.",
    ),
    "S4_state_bond_dimension.png": (
        "Supplementary Figure S4. Bond dimension required by the evolved state, by family and "
        "size.",
        "Largest bond dimension needed to reach cosine similarity 0.999 against the exact "
        "reference, per landscape family, on a logarithmic axis of base two. The requirement "
        "remains small across families and sizes, including for the two families whose "
        "operator saturates its bond-dimension ceiling, so hardness read from the operator "
        "over-estimates the hardness of representing the state. Values at fourteen loci are "
        "lower bounds where the time allocation stopped the search.",
    ),
}

TABLE_TITLES = {
    "S1": "Supplementary Table S1. Acceptance specification: statistics, thresholds, parameter "
    "grids, seeds and time allocations.",
    "S2": "Supplementary Table S2. Validation suite results, all recorded quantities.",
    "S3": "Supplementary Table S3. Landscape ruggedness statistics by size and connectivity.",
    "S4": "Supplementary Table S4. Threshold location and spectral gap per system size.",
    "S4b": "Supplementary Table S4b. Threshold-location criterion, summary statistics.",
    "S5": "Supplementary Table S5. Resource model for the two quantum routes.",
    "S6": "Supplementary Table S6. Time-allocation exclusions and the sensitivity analysis.",
    "S7": "Supplementary Table S7. Operator bond dimension and Pauli-term counts.",
    "S8": "Supplementary Table S8. State bond dimension required for cosine similarity 0.999.",
    "S9": "Supplementary Table S9. Failing-condition tally across scored groups.",
    "S10": "Supplementary Table S10. Hardware results per sweep point.",
    "S11": "Supplementary Table S11. Prior-art verification status per entry.",
    "S12": "Supplementary Table S12. Reproduction comparison across committed artefacts.",
}

# Tables too long to typeset in full are shown truncated in the PDF, with the complete content
# in the accompanying CSV. Stating the row count keeps the omission visible.
MAX_PDF_ROWS = 28


def styles():
    sheet = getSampleStyleSheet()
    sheet.add(
        ParagraphStyle(
            "DocTitle",
            parent=sheet["Title"],
            fontName="Times-Bold",
            fontSize=15,
            leading=19,
            alignment=TA_CENTER,
        )
    )
    sheet.add(
        ParagraphStyle(
            "Byline",
            parent=sheet["Normal"],
            fontName="Times-Roman",
            fontSize=11,
            leading=15,
            alignment=TA_CENTER,
        )
    )
    sheet.add(
        ParagraphStyle(
            "ItemTitle",
            parent=sheet["Normal"],
            fontName="Times-Bold",
            fontSize=10.5,
            leading=14,
            spaceBefore=6,
            spaceAfter=4,
        )
    )
    sheet.add(
        ParagraphStyle(
            "Body",
            parent=sheet["Normal"],
            fontName="Times-Roman",
            fontSize=10,
            leading=14,
            spaceAfter=6,
        )
    )
    sheet.add(
        ParagraphStyle("Mono", parent=sheet["Normal"], fontName="Courier", fontSize=7.2, leading=9)
    )
    return sheet


def scaled_image(path: pathlib.Path, width: float) -> Image:
    """Place a figure at a fixed width, preserving its aspect ratio.

    reportlab will not infer one dimension from the other, so the ratio is read from the
    file. Passing a height of None silently produces an unusable flowable rather than an
    error at the call site.
    """
    from PIL import Image as PILImage

    with PILImage.open(path) as handle:
        source_width, source_height = handle.size
    return Image(str(path), width=width, height=width * source_height / source_width)


def escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def read_csv(path: pathlib.Path) -> tuple[list[str], list[list[str]]]:
    with path.open(newline="", encoding="utf-8") as handle:
        rows = list(csv.reader(handle))
    return (rows[0], rows[1:]) if rows else ([], [])


def csv_flowables(
    path: pathlib.Path, sheet, max_rows: int = MAX_PDF_ROWS, available: float = 163 * mm
):
    header, rows = read_csv(path)
    if not header:
        return [Paragraph("(empty table)", sheet["Body"])]
    shown = rows[:max_rows]
    data = [[Paragraph(f"<b>{escape(str(c))}</b>", sheet["Mono"]) for c in header]]
    for row in shown:
        data.append([Paragraph(escape(str(c)), sheet["Mono"]) for c in row])
    # Column widths are fixed rather than left to reportlab. Without them a single long
    # cell, such as a full acceptance-condition sentence, is laid out on one line and the
    # table is reported as too tall for the frame instead of wrapping.
    columns = len(header)
    widths = [available / columns] * columns
    table = Table(data, colWidths=widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#999999")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EDEDED")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]
        )
    )
    out = [table]
    if len(rows) > max_rows:
        out.append(Spacer(1, 3 * mm))
        out.append(
            Paragraph(
                f"First {max_rows} of {len(rows)} rows shown. The complete table is provided as "
                f"{path.name} with the submission.",
                sheet["Body"],
            )
        )
    return out


def markdown_paragraphs(text: str, sheet) -> list:
    """Render the note and README sources without a Markdown engine.

    Only the constructs these sources actually use are handled: headings, bold spans, bullets,
    and paragraphs. Anything more would be a Markdown implementation, which is not wanted here.
    """
    out = []
    for block in re.split(r"\n\s*\n", text):
        block = block.strip()
        if not block or block.startswith("---"):
            continue
        if block.startswith("#"):
            level = len(block) - len(block.lstrip("#"))
            heading = escape(block.lstrip("#").strip())
            style = "ItemTitle" if level >= 3 else "ItemTitle"
            out.append(Spacer(1, 4 * mm))
            out.append(Paragraph(f"<b>{heading}</b>", sheet[style]))
            continue
        if block.lstrip().startswith(("- ", "| ")):
            for line in block.splitlines():
                line = line.strip()
                if line.startswith("|"):
                    continue
                if line.startswith("- "):
                    body = escape(line[2:])
                    body = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", body)
                    out.append(Paragraph(f"&bull; {body}", sheet["Body"]))
            continue
        body = escape(" ".join(block.split()))
        body = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", body)
        body = re.sub(r"`(.+?)`", r"<font face='Courier'>\1</font>", body)
        out.append(Paragraph(body, sheet["Body"]))
    return out


def build_supplementary_pdf(sources: pathlib.Path, out: pathlib.Path) -> str:
    sheet = styles()
    target = out / "Supplementary_Information.pdf"
    document = SimpleDocTemplate(
        str(target),
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=18 * mm,
        title="Supplementary Information",
        author=AUTHORS,
    )
    story: list = [
        Spacer(1, 30 * mm),
        Paragraph("Supplementary Information", sheet["DocTitle"]),
        Spacer(1, 8 * mm),
        Paragraph(escape(TITLE), sheet["DocTitle"]),
        Spacer(1, 6 * mm),
        Paragraph(escape(AUTHORS), sheet["Byline"]),
        Spacer(1, 3 * mm),
        Paragraph(escape(AFFILIATION), sheet["Byline"]),
        Spacer(1, 12 * mm),
        Paragraph(
            "This file contains Supplementary Figures S1 to S4, Supplementary Tables S1 to S12 "
            "and four Supplementary Notes. Supplementary Data 1 and 2 are provided as separate "
            "comma-separated files, as are the complete versions of the longer tables.",
            sheet["Body"],
        ),
        PageBreak(),
    ]

    figures = sources / "Figures"
    for name, (title, legend) in FIGURE_LEGENDS.items():
        path = figures / name
        if not path.is_file():
            continue
        story.append(Paragraph(escape(title), sheet["ItemTitle"]))
        story.append(Spacer(1, 2 * mm))
        story.append(scaled_image(path, 165 * mm))
        story.append(Spacer(1, 3 * mm))
        story.append(Paragraph(escape(legend), sheet["Body"]))
        story.append(PageBreak())

    tables = sources / "Tables_and_Data"
    order = ["S1", "S2", "S3", "S4", "S4b", "S5", "S6", "S7", "S8", "S9", "S10", "S11", "S12"]
    for key in order:
        matches = sorted(tables.glob(f"Supplementary_Table_{key}_*"))
        if not matches:
            continue
        path = matches[0]
        story.append(Paragraph(escape(TABLE_TITLES.get(key, path.stem)), sheet["ItemTitle"]))
        story.append(Spacer(1, 2 * mm))
        if path.suffix == ".csv":
            story.extend(csv_flowables(path, sheet))
        elif path.suffix == ".md":
            story.append(
                Paragraph(
                    "The acceptance specification is reproduced in full as "
                    f"{path.name} with the submission. It states, for every gate, the statistic "
                    "judged, the threshold applied, the parameter grid, the random seeds and the "
                    "time allocation.",
                    sheet["Body"],
                )
            )
        else:
            text = path.read_text(encoding="utf-8")
            for line in text.splitlines()[:40]:
                story.append(Paragraph(escape(line) or "&nbsp;", sheet["Mono"]))
        story.append(PageBreak())

    notes = sources / "Supplementary_Notes.md"
    if notes.is_file():
        text = notes.read_text(encoding="utf-8")
        text = text.split("---", 2)[-1] if text.count("---") >= 2 else text
        story.append(Paragraph("Supplementary Notes", sheet["ItemTitle"]))
        story.extend(markdown_paragraphs(text, sheet))

    document.build(story)
    return f"Supplementary_Information.pdf ({target.stat().st_size / 1e6:.2f} MB)"


def csv_to_pdf(path: pathlib.Path, out: pathlib.Path, title: str) -> str:
    sheet = styles()
    target = out / f"{path.stem}.pdf"
    document = SimpleDocTemplate(
        str(target),
        pagesize=A4,
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
        title=title,
    )
    header, rows = read_csv(path)
    story = [
        Paragraph(escape(title), sheet["ItemTitle"]),
        Spacer(1, 3 * mm),
        Paragraph(f"{len(rows)} rows, {len(header)} columns.", sheet["Body"]),
        Spacer(1, 3 * mm),
    ]
    story.extend(csv_flowables(path, sheet, max_rows=10_000, available=180 * mm))
    document.build(story)
    return f"{target.name} ({target.stat().st_size / 1e6:.2f} MB)"


def markdown_to_docx(path: pathlib.Path, out: pathlib.Path, title: str) -> str:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(14)

    for line in (TITLE, AUTHORS):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(line).italic = line == AUTHORS

    document.add_paragraph()
    for block in re.split(r"\n\s*\n", path.read_text(encoding="utf-8")):
        block = block.strip()
        if not block or block.startswith("---"):
            continue
        if block.startswith("#"):
            level = min(len(block) - len(block.lstrip("#")), 4)
            document.add_heading(block.lstrip("#").strip(), level=level)
            continue
        if block.lstrip().startswith("| "):
            rows = [r for r in block.splitlines() if r.strip().startswith("|")]
            cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
            cells = [c for c in cells if not all(set(x) <= set("-: ") for x in c)]
            if cells:
                table = document.add_table(rows=0, cols=len(cells[0]))
                table.style = "Table Grid"
                for record in cells:
                    row = table.add_row().cells
                    for index, value in enumerate(record[: len(cells[0])]):
                        row[index].text = re.sub(r"[`*]", "", value)
            continue
        if block.lstrip().startswith("- "):
            for line in block.splitlines():
                if line.strip().startswith("- "):
                    document.add_paragraph(
                        re.sub(r"[`*]", "", line.strip()[2:]), style="List Bullet"
                    )
            continue
        document.add_paragraph(re.sub(r"[`*]", "", " ".join(block.split())))

    target = out / f"{path.stem}.docx"
    document.save(str(target))
    return f"{target.name}"


def build_manuscript_docx(source: pathlib.Path, out: pathlib.Path) -> str:
    """Render the manuscript as Word, which is the journal's preferred format.

    Single-column, unjustified, Times New Roman, with page numbers in the footer, following the
    instructions for revised manuscripts. Figures are supplied as separate files rather than
    embedded, so only their legends appear here.
    """
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 2.0

    section = document.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    field_begin = OxmlElement("w:field_begin")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:field_begin")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instr)
    run._r.append(end)

    text = source.read_text(encoding="utf-8")
    for block in re.split(r"\n\s*\n", text):
        block = block.strip()
        if not block or set(block) <= {"-"}:
            continue
        if block.startswith("$$"):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(block.strip("$ \n")).italic = True
            continue
        if block.startswith("#"):
            level = min(len(block) - len(block.lstrip("#")), 4)
            heading = block.lstrip("#").strip()
            if level == 1:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(heading)
                run.bold = True
                run.font.size = Pt(14)
            else:
                document.add_heading(heading, level=level)
            continue
        if block.lstrip().startswith("|"):
            rows = [r for r in block.splitlines() if r.strip().startswith("|")]
            cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
            cells = [c for c in cells if not all(set(x) <= set("-: ") for x in c)]
            if cells:
                table = document.add_table(rows=0, cols=len(cells[0]))
                table.style = "Table Grid"
                for record in cells:
                    row = table.add_row().cells
                    for index, value in enumerate(record[: len(cells[0])]):
                        row[index].text = re.sub(r"[`*]", "", value)
                document.add_paragraph()
            continue
        if re.match(r"^\d+\. ", block):
            for line in block.splitlines():
                document.add_paragraph(re.sub(r"[`*]", "", line.strip()))
            continue
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for piece, bold in _bold_spans(" ".join(block.split())):
            run = paragraph.add_run(re.sub(r"[`]", "", piece))
            run.bold = bold

    target = out / "QUASAR_manuscript.docx"
    document.save(str(target))
    return f"{target.name}"


def _bold_spans(text: str):
    for index, piece in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if piece:
            yield piece, index % 2 == 1


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    root = pathlib.Path(__file__).resolve().parent.parent.parent / "Submission_package"
    parser.add_argument("--sources", default=str(root / "Supplementary_Information"))
    parser.add_argument("--out", default=str(root / "Supplementary_Information"))
    parser.add_argument("--manuscript", default=str(root / "QUASAR_SciRep_manuscript.md"))
    arguments = parser.parse_args()
    sources = pathlib.Path(arguments.sources)
    out = pathlib.Path(arguments.out)
    out.mkdir(parents=True, exist_ok=True)

    print(build_supplementary_pdf(sources, out))

    pdf_dir = out / "Tables_and_Data_PDF"
    pdf_dir.mkdir(exist_ok=True)
    for path in sorted((sources / "Tables_and_Data").glob("*.csv")):
        key = re.search(r"_(S\d+b?|Data_\d)_", path.name)
        label = key.group(1) if key else path.stem
        title = TABLE_TITLES.get(label, f"Supplementary {label.replace('_', ' ')}. {path.stem}")
        print("  " + csv_to_pdf(path, pdf_dir, title))

    manuscript = pathlib.Path(arguments.manuscript)
    docx_dir = out / "Word"
    docx_dir.mkdir(exist_ok=True)
    if manuscript.is_file():
        print("  " + build_manuscript_docx(manuscript, docx_dir))
    docx_dir.mkdir(exist_ok=True)
    for name, title in (
        ("Supplementary_Notes.md", "Supplementary Notes"),
        ("README.md", "Supplementary Information: contents"),
    ):
        path = sources / name
        if path.is_file():
            print("  " + markdown_to_docx(path, docx_dir, title))
    return 0


if __name__ == "__main__":
    sys.exit(main())
